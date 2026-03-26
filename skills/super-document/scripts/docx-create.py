#!/usr/bin/env python3
"""
Word 文档创建工具
使用 python-docx 库创建和编辑 DOCX 文件
"""

import argparse
import sys
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
except ImportError:
    print("请安装 python-docx: pip install python-docx")
    sys.exit(1)


def create_document(template: str = None) -> Document:
    """创建或加载文档"""
    if template and Path(template).exists():
        return Document(template)
    return Document()


def set_font(run, font_name: str = "微软雅黑", size: int = 11):
    """设置字体"""
    run.font.size = Pt(size)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def add_heading(doc: Document, text: str, level: int = 1):
    """添加标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        set_font(run, "微软雅黑", 16 if level == 1 else 14)
    return heading


def add_paragraph(doc: Document, text: str, bold: bool = False, 
                  align: str = "left", indent: float = 0):
    """添加段落"""
    p = doc.add_paragraph()
    
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    if indent:
        p.paragraph_format.first_line_indent = Inches(indent)
    
    run = p.add_run(text)
    run.bold = bold
    set_font(run)
    return p


def add_table(doc: Document, data: list, headers: list = None, 
              style: str = "Table Grid"):
    """添加表格"""
    if not data:
        return None
    
    rows = len(data) + (1 if headers else 0)
    cols = len(headers) if headers else len(data[0])
    
    table = doc.add_table(rows=rows, cols=cols, style=style)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 添加表头
    if headers:
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    set_font(run, size=10)
    
    # 添加数据
    for row_idx, row_data in enumerate(data):
        for col_idx, value in enumerate(row_data):
            cell = table.rows[row_idx + (1 if headers else 0)].cells[col_idx]
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_font(run, size=10)
    
    return table


def add_bullet_list(doc: Document, items: list, numbered: bool = False):
    """添加列表"""
    for item in items:
        if numbered:
            p = doc.add_paragraph(item, style="List Number")
        else:
            p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            set_font(run)
    return


def add_image(doc: Document, image_path: str, width: float = None, 
              height: float = None):
    """添加图片"""
    if not Path(image_path).exists():
        print(f"图片不存在: {image_path}")
        return None
    
    if width and height:
        doc.add_picture(image_path, width=Inches(width), height=Inches(height))
    elif width:
        doc.add_picture(image_path, width=Inches(width))
    elif height:
        doc.add_picture(image_path, height=Inches(height))
    else:
        doc.add_picture(image_path)
    
    return doc.paragraphs[-1]


def add_page_break(doc: Document):
    """添加分页符"""
    doc.add_page_break()


def add_date(doc: Document, format: str = "%Y年%m月%d日"):
    """添加日期"""
    date_str = datetime.now().strftime(format)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(date_str)
    set_font(run, size=10)
    return p


def create_from_markdown(doc: Document, md_content: str):
    """从 Markdown 创建文档"""
    lines = md_content.split("\n")
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # 标题
        if line.startswith("# "):
            add_heading(doc, line[2:], level=1)
        elif line.startswith("## "):
            add_heading(doc, line[3:], level=2)
        elif line.startswith("### "):
            add_heading(doc, line[4:], level=3)
        # 列表
        elif line.startswith("- ") or line.startswith("* "):
            add_bullet_list(doc, [line[2:]])
        elif line.startswith("1. "):
            add_bullet_list(doc, [line[3:]], numbered=True)
        # 空行
        elif not line.strip():
            pass
        # 普通段落
        else:
            add_paragraph(doc, line)
        
        i += 1


def create_report(template: str = None, title: str = None, 
                  content: list = None, output: str = None):
    """创建报告文档"""
    doc = create_document(template)
    
    # 标题
    if title:
        add_heading(doc, title, level=1)
        add_date(doc)
        add_paragraph(doc, "")
    
    # 内容
    if content:
        for item in content:
            if isinstance(item, dict):
                item_type = item.get("type", "paragraph")
                
                if item_type == "heading":
                    add_heading(doc, item["text"], item.get("level", 2))
                elif item_type == "paragraph":
                    add_paragraph(doc, item["text"])
                elif item_type == "table":
                    add_table(doc, item["data"], item.get("headers"))
                elif item_type == "image":
                    add_image(doc, item["path"], item.get("width"))
                elif item_type == "list":
                    add_bullet_list(doc, item["items"], item.get("numbered", False))
                elif item_type == "page_break":
                    add_page_break(doc)
            else:
                add_paragraph(doc, str(item))
    
    # 保存
    if output:
        doc.save(output)
        print(f"✓ 文档已保存到: {output}")
    
    return doc


def main():
    parser = argparse.ArgumentParser(description="Word 文档创建工具")
    parser.add_argument("output", help="输出文件路径")
    parser.add_argument("-t", "--template", help="模板文件路径")
    parser.add_argument("--title", help="文档标题")
    parser.add_argument("-m", "--markdown", help="从 Markdown 文件创建")
    parser.add_argument("--demo", action="store_true", help="创建示例文档")
    
    args = parser.parse_args()
    
    if args.demo:
        # 创建演示文档
        doc = create_document(args.template)
        
        add_heading(doc, "示例文档", level=1)
        add_date(doc)
        add_paragraph(doc, "这是一个使用 python-docx 创建的示例文档。")
        
        add_heading(doc, "1. 文本格式", level=2)
        add_paragraph(doc, "这是普通文本。", bold=False)
        add_paragraph(doc, "这是加粗文本。", bold=True)
        add_paragraph(doc, "这是居中文本。", align="center")
        
        add_heading(doc, "2. 列表", level=2)
        add_bullet_list(doc, ["第一项", "第二项", "第三项"])
        
        add_heading(doc, "3. 表格", level=2)
        add_table(doc, [
            ["张三", 25, "北京"],
            ["李四", 30, "上海"],
            ["王五", 28, "广州"]
        ], headers=["姓名", "年龄", "城市"])
        
        add_heading(doc, "4. 分页示例", level=2)
        add_paragraph(doc, "下一页将展示更多内容...")
        add_page_break(doc)
        
        add_heading(doc, "第二页", level=1)
        add_paragraph(doc, "这是第二页的内容。")
        
        doc.save(args.output)
        print(f"✓ 示例文档已保存到: {args.output}")
    
    elif args.markdown:
        # 从 Markdown 创建
        md_content = Path(args.markdown).read_text(encoding="utf-8")
        doc = create_document(args.template)
        create_from_markdown(doc, md_content)
        doc.save(args.output)
        print(f"✓ 文档已保存到: {args.output}")
    
    elif args.title:
        # 创建报告
        create_report(args.template, args.title, output=args.output)
    
    else:
        # 空文档
        doc = create_document(args.template)
        doc.save(args.output)
        print(f"✓ 空文档已保存到: {args.output}")


if __name__ == "__main__":
    main()